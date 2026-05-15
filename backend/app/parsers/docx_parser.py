from docx import Document
from app.parsers.base import BaseParser, ParsedDocument, ParsedPage
from loguru import logger

class DOCXParser(BaseParser):
    def parse(self, file_path: str, file_id: str) -> ParsedDocument:
        logger.info(f"Parsing DOCX: {file_path}")
        doc = Document(file_path)
        
        # DOCX doesn't have native "pages" in the same way PDF does
        # We'll treat the whole doc as one page or split by something else
        # For simplicity, treat as one page for now, or split by paragraphs
        
        full_text = "\n".join([para.text for para in doc.paragraphs])
        
        pages = [ParsedPage(
            page_number=1,
            text=full_text
        )]
            
        return ParsedDocument(
            file_id=file_id,
            pages=pages,
            metadata={}
        )
